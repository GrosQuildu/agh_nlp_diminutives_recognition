# !/usr/bin/env python
# -*- coding: utf-8 -*-

"""Tool for highlighting diminutives in a docx file.
    Example:
        
    Authors:
    * Izabela Stechnij
    * Dominik Sepioło
    * Paweł Płatek
"""
# fmt: on

import argparse
from copy import copy
from os.path import isfile
from sys import exit

from docx import Document  # type: ignore
from docx.dml.color import ColorFormat  # type: ignore
from docx.enum.dml import MSO_COLOR_TYPE  # type: ignore
from docx.enum.text import WD_COLOR_INDEX  # type: ignore
from docx.exceptions import PythonDocxError  # type: ignore
from docx.text.font import Font  # type: ignore
from docx.text.run import Run  # type: ignore

from rozpoznawaczek import L, find_diminutives


def copy_style(new_element, original_element):
    if type(new_element) != type(original_element):
        raise Exception('Types of objects to copy style do not match')

    to_copy, to_deepcopy = [], []

    if type(new_element) == Run:
        to_copy = ['bold', 'italic', 'underline']
        to_deepcopy = ['font']

    elif type(new_element) == Font:
        to_copy = ['all_caps', 'bold', 'cs_bold', 'cs_italic',
                   'double_strike', 'emboss', 'hidden', 'highlight_color',
                   'imprint', 'italic', 'name', 'outline',
                   'rtl', 'shadow', 'size', 'small_caps', 'snap_to_grid', 'spec_vanish',
                   'strike', 'subscript', 'superscript', 'underline', 'web_hidden'
                   ]
        to_deepcopy = ['color']

    elif type(new_element) == ColorFormat:
        if original_element.type == MSO_COLOR_TYPE.THEME:
            to_copy = ['theme_color']
        elif original_element.type == MSO_COLOR_TYPE.RGB:
            to_copy = ['rgb']

    else:
        return

    for attr in to_copy:
        if hasattr(original_element, attr):
            setattr(new_element, attr, getattr(original_element, attr))

    for attr in to_deepcopy:
        if hasattr(new_element, attr) and hasattr(original_element, attr):
            copy_style(getattr(new_element, attr),
                       getattr(original_element, attr))


def highlight(document, color):
    for paragraph in document.paragraphs:
        runs = copy(paragraph.runs)
        paragraph.clear()

        for run in runs:
            original_text = run.text
            diminutives = find_diminutives(original_text)

            coursor = 0
            for start_position, end_position in diminutives:
                # normal text
                new_run = paragraph.add_run(
                    original_text[coursor:start_position], run.style)
                copy_style(new_run, run)

                # diminutive
                new_run = paragraph.add_run(
                    original_text[start_position:end_position], run.style)
                copy_style(new_run, run)
                new_run.font.highlight_color = color

                coursor = end_position

            # normal text
            new_run = paragraph.add_run(
                original_text[coursor:len(original_text)], run.style)
            copy_style(new_run, run)

    return document


def main():
    colors = [attr for attr in dir(
        WD_COLOR_INDEX) if attr.isupper() and not attr.startswith('_')]
    default_color = 'RED' if 'RED' in colors else colors[-1]

    parser = argparse.ArgumentParser(description='Hightlight diminutives')
    parser.add_argument(
        '-i', '--input', help='Input docx file', required=True)
    parser.add_argument(
        '-o', '--output', help='Output docx file', required=True)
    parser.add_argument(
        '-f', '--force', help='Force output overwrite', action='store_true')
    parser.add_argument('-c', '--color', help='Color to highlight diminutives',
                        choices=colors, default=default_color)
    parser.add_argument("-v", "--verbose", help="debug output",
                        action="store_true")

    args = parser.parse_args()

    L.setLevel('INFO')
    if args.verbose:
        L.setLevel('DEBUG')

    if not hasattr(WD_COLOR_INDEX, args.color):
        L.error('WD_COLOR_INDEX has no color `%s`', args.color)
        return 1

    if not isfile(args.input):
        L.error('Not such file: %s', args.input)
        return 1

    if not args.force and isfile(args.output):
        L.error('File exists: %s. Use -f/--force to overwrite.', args.output)
        return 1

    try:
        document = Document(args.input)
    except PythonDocxError as e:
        L.error('Error when opening input file: %s', e)
        return 1

    highlighted_document = highlight(document, getattr(WD_COLOR_INDEX, args.color))
    highlighted_document.save(args.output)


if __name__ == "__main__":
    exit(main())
